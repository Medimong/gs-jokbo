# -*- coding: utf-8 -*-
"""외과 족보 docx 파서.

본문을 문단/표/이미지 순서 그대로 순회한 뒤 문항 단위로 자른다.
정답은 ①"답: 3번" 텍스트 ②선지 런의 서식(빨강·형광·볼드) 두 경로로 잡는다.
"""
import os, re, sys, json, hashlib, unicodedata
from docx import Document
from docx.shared import RGBColor

NS = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
      'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
      'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

SECTIONS = {  # 원본 표기 -> 정규 분과 코드
    'LV': 'LV', '간': 'LV', 'HBP': 'BP', 'BP': 'BP', 'HPB': 'BP',
    'BR': 'BR', 'BREAST': 'BR', 'ENDO': 'ENDO', 'END': 'ENDO',
    'UGI': 'UGI', 'LGI': 'CR', 'CR': 'CR', 'CRS': 'CR', 'COLORECTAL': 'CR',
    'GI': 'GI', 'VS': 'VS', 'VAS': 'VS', 'VASCULAR': 'VS',
    'PDS': 'PDS', 'PS': 'PDS', 'PED': 'PDS',
    'TACS': 'TACS', 'TRS': 'TACS', 'TRAUMA': 'TACS',
    'TX': 'TX', '총론': '총론', '해부학': '총론',
}
SEC_RE = re.compile(r'^\s*\[?([A-Za-z가-힣]{2,12})\]?\s*[\(（]?\s*(\d{1,2})?\s*[\)）]?\s*$')
Q_NUM_RE = re.compile(r'^\s*(\d{1,3})\s*[.)]\s*(\S.*)$')
Q_TAG_RE = re.compile(r'^\s*(\((?:\d{2}[-–]\d?[A-D]{1,2}|\d{4}|18이전|1[0-9]이전)[^)]*\)\s*)+')
Q_BRACKET_RE = re.compile(r'^\s*[<〈]\s*(\d{1,3})\s*([^>〉]*)[>〉]\s*(.*)$')
OPT_RE = re.compile(r'^\s*(?:([1-5])\s*[.)]|([①-⑤])|([a-eA-E])\s*[.)])\s*(\S.*)$')
CIRCLED = '①②③④⑤'
ANS_RE = re.compile(r'^\s*(?:정답|답|answer)\s*[:：)\]]?\s*(.*)$', re.I)
ANS_NUM_RE = re.compile(r'([1-5])\s*번|^\s*([1-5])\s*$|([①-⑤])')

def nfc(s): return unicodedata.normalize('NFC', s or '')

def run_marked(run):
    """정답 표시 서식인가: 빨강 계열 글자색 / 형광 / (볼드는 약한 신호)"""
    f = run.font
    try:
        if f.highlight_color is not None: return 'highlight'
    except Exception: pass
    try:
        c = f.color
        if c is not None and c.rgb is not None:
            v = int(str(c.rgb), 16)
            r, g, b = (v >> 16) & 255, (v >> 8) & 255, v & 255
            if r > 120 and r > g + 60 and r > b + 60: return 'red'
            if b > 120 and b > r + 60 and b > g + 60: return 'blue'
    except Exception: pass
    if f.bold: return 'bold'
    return None

def iter_body(doc):
    """문단·표·이미지를 문서에 나온 순서 그대로 산출."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split('}')[1]
        if tag == 'p':
            blips = child.findall('.//a:blip', NS)
            para = None
            for p in doc.paragraphs:
                if p._p is child: para = p; break
            text = nfc(para.text.strip()) if para is not None else ''
            marks = []
            if para is not None:
                for run in para.runs:
                    if run.text.strip():
                        marks.append((nfc(run.text), run_marked(run)))
            if text:
                yield {'t': 'p', 'text': text, 'marks': marks}
            for b in blips:
                rid = b.get('{%s}embed' % NS['r'])
                if rid: yield {'t': 'img', 'rid': rid}
        elif tag == 'tbl':
            for tb in doc.tables:
                if tb._tbl is child:
                    rows = [[nfc(c.text.strip()) for c in r.cells] for r in tb.rows]
                    yield {'t': 'tbl', 'rows': rows}
                    break
            for b in child.findall('.//a:blip', NS):
                rid = b.get('{%s}embed' % NS['r'])
                if rid: yield {'t': 'img', 'rid': rid}

def norm_section(label):
    key = re.sub(r'[\[\]()（）\s]', '', label).upper()
    return SECTIONS.get(key) or SECTIONS.get(nfc(label).strip())

def parse(path, rel):
    doc = Document(path)
    items = list(iter_body(doc))
    rid_blob = {}
    for rid, part in doc.part.related_parts.items():
        if 'image' in str(getattr(part, 'content_type', '')):
            try: rid_blob[rid] = part.blob
            except Exception: pass
    return parse_items(items, rel, rid_blob), rid_blob


def parse_items(items, rel, rid_blob=None):
    """문단/표/이미지 스트림을 문항 단위로 자른다. docx·pdf 공용."""
    rid_blob = rid_blob or {}
    # 문항 번호가 거의 없는 복원본은 '정답 줄'을 문항 경계로 삼는 모드로 전환한다
    numbered = sum(1 for it in items if it['t'] == 'p'
                   and (m := Q_NUM_RE.match(it['text'])) and len(m.group(2)) > 12)
    numberless = numbered < 5
    qs, cur, section, topic = [], None, None, None
    last_num = 0
    pending_head = None   # <1 주제> 처럼 발문이 다음 문단에 오는 경우

    def flush():
        nonlocal cur
        if cur and (cur['options'] or len(cur['stem']) > 20):
            qs.append(cur)
        cur = None

    for it in items:
        if it['t'] == 'img':
            if cur is not None: cur['img_rids'].append(it['rid'])
            continue
        if it['t'] == 'tbl':
            if cur is not None: cur['tables'].append(it['rows'])
            continue
        text, marks = it['text'], it['marks']

        m = SEC_RE.match(text)
        if m and len(text) <= 24:
            code = norm_section(m.group(1))
            if code:
                flush(); section = code
                topic = None if not text.startswith('[') else nfc(m.group(1))
                continue
        if text.startswith('[') and text.endswith(']') and len(text) <= 24:
            code = norm_section(text)
            if code: flush(); section = code
            else: topic = text.strip('[]')
            continue

        am = ANS_RE.match(text)
        if am and cur is not None and len(text) < 200:
            cur['answer_raw'] = text
            continue

        om = OPT_RE.match(text)
        qm = Q_NUM_RE.match(text)
        tagm = Q_TAG_RE.match(text)
        bm = Q_BRACKET_RE.match(text)

        if bm:
            flush()
            last_num = int(bm.group(1))
            pending_head = {'num': last_num, 'topic': bm.group(2).strip() or None}
            rest = bm.group(3).strip()
            cur = {'src': rel, 'section': section, 'topic': pending_head['topic'] or topic,
                   'num': last_num, 'stem': rest, 'options': [], 'opt_marks': [],
                   'answer_raw': None, 'img_rids': [], 'tables': [], 'src_tags': [],
                   'extra': []}
            continue

        # 문항 시작: 번호+긴 발문 (선지 번호와 구분: 발문은 길고 뒤에 선지가 따라옴)
        is_new_q = False
        if tagm and len(text) > 25:
            is_new_q = True
        elif qm and len(qm.group(2)) > 12:
            n = int(qm.group(1))
            # 선지/정답을 이미 채운 뒤이거나, 번호가 직전 문항 다음으로 이어지면 새 문항
            if cur is None or cur['options'] or cur.get('answer_raw') or n == last_num + 1:
                is_new_q = True

        # 번호 없는 복원본: 정답 줄 뒤에 오는 일반 텍스트 = 다음 문항 발문
        if (numberless and not is_new_q and cur is not None and cur.get('answer_raw')
                and not om and len(text) > 10):
            flush()
            cur = {'src': rel, 'section': section, 'topic': topic, 'num': None,
                   'stem': text, 'options': [], 'opt_marks': [], 'answer_raw': None,
                   'img_rids': [], 'tables': [], 'src_tags': [], 'extra': []}
            continue

        # 번호 없는 복원본의 첫 문항: 분과 헤딩 뒤 첫 긴 줄
        if (numberless and cur is None and section is not None
                and not om and len(text) > 15):
            cur = {'src': rel, 'section': section, 'topic': topic, 'num': None,
                   'stem': text, 'options': [], 'opt_marks': [], 'answer_raw': None,
                   'img_rids': [], 'tables': [], 'src_tags': [], 'extra': []}
            continue

        if is_new_q:
            flush()
            num = int(qm.group(1)) if qm else None
            if num: last_num = num
            body = qm.group(2) if qm else text[tagm.end():].strip()
            tags = re.findall(r'\(([^)]*)\)', tagm.group(0)) if tagm else []
            cur = {'src': rel, 'section': section, 'topic': topic, 'num': num,
                   'stem': body, 'options': [], 'opt_marks': [], 'answer_raw': None,
                   'img_rids': [], 'tables': [], 'src_tags': tags, 'extra': []}
            continue

        if cur is None:
            continue

        if not cur['stem'] and not om and len(text) > 12:
            cur['stem'] = text
            continue

        if om:
            idx = om.group(1) or om.group(3)
            if om.group(2): idx = str(CIRCLED.index(om.group(2)) + 1)
            body = om.group(4)
            cur['options'].append(body)
            cur['opt_marks'].append(any(k in ('red', 'highlight') for _, k in marks))
            continue

        # 번호 없는 선지 후보 = 짧은 줄. 일단 extra로 모으고 후처리에서 판정.
        cur['extra'].append({'text': text,
                             'marked': any(k in ('red', 'highlight') for _, k in marks)})
    flush()

    # 후처리: 번호 없는 선지 복구 + 정답 확정 + 이미지 저장
    for q in qs:
        if not q['options']:
            short = [e for e in q['extra'] if len(e['text']) <= 70]
            if len(short) >= 3 and len(short) == len(q['extra']):
                q['options'] = [e['text'] for e in short]
                q['opt_marks'] = [e['marked'] for e in short]
                q['extra'] = []
            elif len(short) >= 3:
                tail = []
                for e in reversed(q['extra']):
                    if len(e['text']) <= 70: tail.append(e)
                    else: break
                if 3 <= len(tail) <= 6:
                    tail.reverse()
                    q['options'] = [e['text'] for e in tail]
                    q['opt_marks'] = [e['marked'] for e in tail]
                    q['extra'] = q['extra'][:-len(tail)]
        if q['extra']:
            q['stem'] = (q['stem'] + '\n' + '\n'.join(e['text'] for e in q['extra'])).strip()
        q.pop('extra', None)

        ans, ans_text = None, None
        if q['answer_raw']:
            tail = ANS_RE.match(q['answer_raw']).group(1).strip()
            m = ANS_NUM_RE.search(tail)
            if m:
                g = m.group(1) or m.group(2)
                ans = int(g) if g else CIRCLED.index(m.group(3)) + 1
            if tail and (ans is None or len(tail) > 4):
                ans_text = tail
        q['answer_text'] = ans_text
        if ans is None and any(q['opt_marks']):
            marked = [i + 1 for i, v in enumerate(q['opt_marks']) if v]
            if len(marked) == 1: ans = marked[0]
        q['answer'] = ans
        q['answer_src'] = ('text' if q['answer_raw'] and ans else
                           'format' if ans else None)
        q['images'] = []
        for rid in q['img_rids']:
            blob = rid_blob.get(rid)
            if blob:
                q['images'].append({'md5': hashlib.md5(blob).hexdigest()[:10],
                                    'bytes': len(blob)})
            else:
                q['images'].append({'md5': rid, 'bytes': None})
        q.pop('img_rids', None); q.pop('opt_marks', None)
    return qs

if __name__ == '__main__':
    inv = json.load(open(os.path.join(os.path.dirname(__file__), 'a0_files.json')))
    rel = sys.argv[1]
    qs, _ = parse(os.path.join(inv['drive'], rel), rel)
    print(f"{rel}: {len(qs)} questions")
    for q in qs[:int(sys.argv[2]) if len(sys.argv) > 2 else 3]:
        print(json.dumps(q, ensure_ascii=False, indent=1)[:1400])
